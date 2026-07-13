from datetime import date
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading_centered(document, text, level=0):
    h = document.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return h


def add_case_table(document, case):
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    rows = [
        ("Test Case ID", case["id"]),
        ("Title", case["title"]),
        ("Objective", case["objective"]),
        ("Preconditions", case["preconditions"]),
        ("Test Data", case["test_data"]),
        ("Priority", case["priority"]),
    ]

    for key, value in rows:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value

    steps_row = table.add_row().cells
    steps_row[0].text = "Test Steps"
    steps_row[1].text = "\n".join([f"{i + 1}. {step}" for i, step in enumerate(case["steps"])])

    expected_row = table.add_row().cells
    expected_row[0].text = "Expected Results"
    expected_row[1].text = "\n".join([f"{i + 1}. {res}" for i, res in enumerate(case["expected_results"])])

    document.add_paragraph("")


def main():
    doc = Document()

    # Base font setup
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    add_heading_centered(doc, "Software Test Case Specification", level=0)
    p = doc.add_paragraph("IEEE 829 Style Format")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    doc.add_paragraph("Project: GigWork Platform Frontend")
    doc.add_paragraph("Document ID: GWF-IEEE-TCS-001")
    doc.add_paragraph(f"Date: {date.today().isoformat()}")
    doc.add_paragraph("Version: 1.0")
    doc.add_paragraph("Prepared By: QA Team")

    doc.add_page_break()

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This document defines executable frontend test cases for authentication, "
        "worker ranking, review submission, and request dispatch workflows. "
        "The structure follows IEEE 829 test case specification style."
    )

    doc.add_heading("2. Scope", level=1)
    doc.add_paragraph(
        "The test cases validate UI behavior, form validation, API-integration flow "
        "through mocks, and user-facing outcomes in the React frontend."
    )

    doc.add_heading("3. Test Environment", level=1)
    doc.add_paragraph("Frontend: React + Vite")
    doc.add_paragraph("Test Framework: Vitest + React Testing Library")
    doc.add_paragraph("Execution Mode: Automated component/integration tests")

    doc.add_heading("4. Test Case Specifications", level=1)

    test_cases = [
        {
            "id": "TC-LOGIN-01",
            "title": "Successful login with valid credentials",
            "objective": "Verify user can log in and proceed to authorized flow.",
            "preconditions": "Registered active user account exists.",
            "test_data": "email=test@example.com, password=Password123",
            "priority": "High",
            "steps": [
                "Open Login page.",
                "Enter valid email and password.",
                "Click Sign In."
            ],
            "expected_results": [
                "Login request is submitted.",
                "Authentication succeeds.",
                "User state transitions to authenticated."
            ]
        },
        {
            "id": "TC-LOGIN-02",
            "title": "Login failure with invalid credentials",
            "objective": "Verify incorrect credentials return a clear error.",
            "preconditions": "Registered user exists.",
            "test_data": "email=test@example.com, password=WrongPass",
            "priority": "High",
            "steps": [
                "Open Login page.",
                "Enter valid email and invalid password.",
                "Click Sign In."
            ],
            "expected_results": [
                "Authentication fails.",
                "Error alert is displayed.",
                "User remains unauthenticated."
            ]
        },
        {
            "id": "TC-SIGNUP-01",
            "title": "Successful signup with valid data",
            "objective": "Verify a new user can register with valid input.",
            "preconditions": "Email is not already registered.",
            "test_data": "Valid first name, last name, email, phone, matching passwords",
            "priority": "High",
            "steps": [
                "Open Register page.",
                "Fill all required fields with valid values.",
                "Click Create Account."
            ],
            "expected_results": [
                "Registration request is submitted.",
                "Success message is shown.",
                "User is redirected to login flow."
            ]
        },
        {
            "id": "TC-SIGNUP-02",
            "title": "Signup validation for mismatched passwords",
            "objective": "Ensure client validation blocks mismatched passwords.",
            "preconditions": "Register page is accessible.",
            "test_data": "password=Password123, password2=Password321",
            "priority": "High",
            "steps": [
                "Open Register page.",
                "Fill required fields with mismatched passwords.",
                "Click Create Account."
            ],
            "expected_results": [
                "Submission is blocked.",
                "Password mismatch message is shown.",
                "No register API call is accepted as success."
            ]
        },
        {
            "id": "TC-RANK-01",
            "title": "Search workers and verify ranking order",
            "objective": "Validate workers are displayed in descending final score.",
            "preconditions": "Customer session available; categories and workers service returns data.",
            "test_data": "category=Electrician, radius=10",
            "priority": "High",
            "steps": [
                "Open Find Workers page.",
                "Select category and keep radius value.",
                "Click Search Workers."
            ],
            "expected_results": [
                "Worker list appears.",
                "Top-ranked worker appears first.",
                "Ranking score and worker details are visible."
            ]
        },
        {
            "id": "TC-RANK-02",
            "title": "No-worker result handling",
            "objective": "Verify empty-state behavior when no workers match criteria.",
            "preconditions": "Customer session available.",
            "test_data": "Category/radius returning no records",
            "priority": "Medium",
            "steps": [
                "Open Find Workers page.",
                "Search with filters that return no workers."
            ],
            "expected_results": [
                "No workers found message is displayed.",
                "UI remains responsive and stable."
            ]
        },
        {
            "id": "TC-REVIEW-01",
            "title": "Successful review and rating submission",
            "objective": "Verify customer can submit rating and text review for a request.",
            "preconditions": "Completed request exists for the customer.",
            "test_data": "request=1, rating=1..5, comment=Great service",
            "priority": "High",
            "steps": [
                "Open Submit Review page with request query parameter.",
                "Select star rating.",
                "Enter review comment.",
                "Click Submit Review."
            ],
            "expected_results": [
                "Review API is called with request, rating, and review text.",
                "Success message appears.",
                "User is redirected to requests view."
            ]
        },
        {
            "id": "TC-REVIEW-02",
            "title": "Prevent review submit without rating",
            "objective": "Ensure rating is mandatory for submission.",
            "preconditions": "Submit Review page loaded.",
            "test_data": "rating=0, comment provided",
            "priority": "High",
            "steps": [
                "Open Submit Review page.",
                "Keep rating at zero.",
                "Attempt to submit review."
            ],
            "expected_results": [
                "Submit action is disabled or blocked.",
                "No review API success path is triggered."
            ]
        },
        {
            "id": "TC-REQ-01",
            "title": "Create service request to selected worker",
            "objective": "Verify request payload includes preferred worker ID from query parameter.",
            "preconditions": "Customer selected worker from search results.",
            "test_data": "worker=9, category=2, title and description valid",
            "priority": "High",
            "steps": [
                "Open Create Request page with worker and category query params.",
                "Enter request title and description.",
                "Click Create Request."
            ],
            "expected_results": [
                "Request API is called with category, title, description, and preferred_worker_id.",
                "Success message is displayed."
            ]
        },
        {
            "id": "TC-REQ-02",
            "title": "Required field validation in create request",
            "objective": "Ensure mandatory request fields must be filled.",
            "preconditions": "Create Request page loaded.",
            "test_data": "Missing title and/or description",
            "priority": "High",
            "steps": [
                "Open Create Request page.",
                "Leave mandatory fields empty.",
                "Click Create Request."
            ],
            "expected_results": [
                "Submission is blocked by validation.",
                "No successful request creation occurs."
            ]
        }
    ]

    for case in test_cases:
        doc.add_heading(case["id"], level=2)
        add_case_table(doc, case)

    doc.add_heading("5. Traceability", level=1)
    doc.add_paragraph(
        "These automated test scripts are implemented in frontend test files under src/pages "
        "using Vitest and React Testing Library."
    )

    output_path = "IEEE_Frontend_Test_Case_Specification.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")


if __name__ == "__main__":
    main()
